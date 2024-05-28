from tkinter import Tk, Text, Button, Label, END, messagebox, Checkbutton, IntVar
from tkinter import filedialog
from docx import Document
from docx.shared import Pt, RGBColor
import subprocess

# Format the input text for bionic reading by splitting each word into bold and normal parts.
def bionic_reading(text):
    lines = text.splitlines()  # Split text into lines
    formatted_lines = []
    for line in lines:
        words = line.split()  # Split line into words
        formatted_words = []
        for word in words:
            if len(word) > 1:
                mid = len(word) // 2  # Calculate the middle index of the word
                formatted_words.append((word[:mid], word[mid:]))  # Split word into two parts
            else:
                formatted_words.append((word, ''))  # Handle single-character words
        formatted_lines.append(formatted_words)
    return formatted_lines

# Apply a color gradient to text.
def apply_gradient(run, start_color, end_color, ratio):
    start_r, start_g, start_b = start_color
    end_r, end_g, end_b = end_color
    new_r = int(start_r + (end_r - start_r) * ratio)
    new_g = int(start_g + (end_g - start_g) * ratio)
    new_b = int(start_b + (end_b - start_b) * ratio)
    run.font.color.rgb = RGBColor(new_r, new_g, new_b)

# Export the formatted text to a Word document.
def export_to_word(formatted_lines, filename, bionic, shading):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Candara'  # Set default font
    font.size = Pt(12)  # Set default font size

    # Define color gradients
    gradients = [
        (RGBColor(0, 0, 0), RGBColor(0, 0, 139)),  # Black to dark blue
        (RGBColor(0, 0, 139), RGBColor(0, 0, 0)),  # Dark blue to black
        (RGBColor(0, 0, 0), RGBColor(128, 0, 128)),  # Black to purple
        (RGBColor(128, 0, 128), RGBColor(0, 0, 0))  # Purple to black
    ]

    for i, formatted_words in enumerate(formatted_lines):
        paragraph = document.add_paragraph()
        start_color, end_color = gradients[i % 4]
        total_length = sum(len(bold_part + normal_part) for bold_part, normal_part in formatted_words)
        current_length = 0

        for bold_part, normal_part in formatted_words:
            word_length = len(bold_part + normal_part)
            ratio = current_length / total_length
            run = paragraph.add_run(bold_part)
            if bionic:
                run.bold = True
            if shading:
                apply_gradient(run, start_color, end_color, ratio)
            run = paragraph.add_run(normal_part)
            run.font.name = 'Candara'
            if shading:
                apply_gradient(run, start_color, end_color, ratio)
            paragraph.add_run(' ')
            current_length += word_length

    document.save(filename)
    messagebox.showinfo("Success", "The document has been created successfully.")
    open_file(filename)

# Import text from a Word document.
def import_word_document():
    filename = filedialog.askopenfilename(filetypes=[("Word Document", "*.docx")], title="Select Word Document")
    if filename:
        document = Document(filename)
        text = ""
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    return None

# Handle the click event for the convert button.
def on_convert_click():
    text = text_box.get("1.0", END).strip()
    if not text:
        text = import_word_document()
        if not text:
            messagebox.showwarning("Warning", "No text or Word document selected.")
            return
    formatted_lines = bionic_reading(text)
    filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], title="Save Document As")
    if filename:
        export_to_word(formatted_lines, filename, bionic_var.get(), shading_var.get())
        root.destroy()  # Close the window after conversion

# Open the generated Word document.
def open_file(filename):
    subprocess.call(["open", filename])

# Create the main window
root = Tk()
root.title("Text Formatter")

# Set window size and position
window_width = 450
window_height = 350
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()
x_coordinate = (screen_width - window_width) // 2
y_coordinate = (screen_height - window_height) // 2
root.geometry("{}x{}+{}+{}".format(window_width, window_height, x_coordinate, y_coordinate))
root.configure(bg='#f0f0f0')

# Set padding for widgets
padding = 10

# Create a Text widget for input
text_box = Text(root, height=8, width=50, font=("Arial", 12), wrap="word", relief="flat", bd=1)
text_box.pack(pady=(padding, 0), padx=padding)

# Create a Label to provide explanation
explanation_label = Label(root, text="Input plain text or import a Word document. Select formatting options below.", font=("Arial", 10), bg='#f0f0f0')
explanation_label.pack(pady=padding)

# Create checkboxes for options
bionic_var = IntVar(value=1)
shading_var = IntVar(value=1)
bionic_checkbox = Checkbutton(root, text="Bionic Reading", variable=bionic_var, font=("Arial", 10), bg='#f0f0f0')
shading_checkbox = Checkbutton(root, text="Paragraph Shading", variable=shading_var, font=("Arial", 10), bg='#f0f0f0')
bionic_checkbox.pack(pady=padding)
shading_checkbox.pack(pady=padding)

# Create a Button to import Word document or trigger conversion
convert_button = Button(root, text="Convert to Word Document", command=on_convert_click, font=("Arial", 12), bg='#007BFF', fg='black', relief="flat", bd=0, padx=10, pady=5)
convert_button.pack(pady=padding)

# Run the application
root.mainloop()
